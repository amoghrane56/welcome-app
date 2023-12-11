import streamlit as st
from docx import Document
import pandas as pd
import pyodbc
from docx.shared import Pt

# Define a function to fetch data from the database
def fetch_data_from_database(pms_account_code):
    # Establish a connection to the SQL Server database
    connection_string = """DRIVER={SQL Server};SERVER=192.168.13.10;
                          DATABASE=IntegraUAT;UID=amog;PWD=Abcd#123"""
    conn = pyodbc.connect(connection_string)
    cursor = conn.cursor()

    # Use parameterized query to avoid SQL injection
    sql_query = """
    SELECT 
        A.client_code,
        A.clientname AS Accountname, 
        C.Clientname, 
        ISNULL(a.address1,'') + ' ' + ISNULL(a.CITY,'') + ' ' + ISNULL(a.State,'') + ' ' + ISNULL(a.pin,'') + ' ' + ISNULL(a.Country,'') AS Address,  
        A.ActiveDate,
        A.backofficecodeequity, 
        D.MainObjective AS SchemeCodename,
        G.TypeDesc AS BenchMark, 
        A.EMAIL,
        A.mobile_no,
        B.Int_Name AS IntroducerDistributorName,
        '' AS FeesCommissionDistributor, 
        A.ctPersonDecision AS RelationshipManager,
        A.ctPersonDEmail AS RMEmail, 
        '' AS MobileNo, 
        E.OpeningEquityCorpus + E.OpeningCashCorpus AS TotalCorpusIntroduced, 
        E.OpeningCashCorpus AS Fund, 
        E.OpeningEquityCorpus AS Securities, 
        A.Usr_clientid AS LoginId, 
        '******' AS Password
    FROM hdr_client A 
        INNER JOIN HDR_Scheme D ON A.SchemeCode = D.SchemeCode  
        INNER JOIN HDR_ClientHead C ON A.head_clientcode = C.Client_code
        INNER JOIN HDR_Intermediary B ON A.inter_code = B.int_code 
        INNER JOIN hist_clientnav E ON A.client_code = E.ClientCode 
        INNER JOIN dtl_schemeportfolio_benchmark_map F ON D.SchemeCode = F.Scheme_Code
        INNER JOIN HDR_SensexType G ON G.TypeCode = F.Benchmarkindices
    WHERE 
        A.backofficecodeequity = ? 
        AND A.SubBrokerCode IS NOT NULL
        AND E.NavAsOn in (select MAX(NavAsOn) from hist_clientnav);
    """

    df = pd.read_sql_query(sql_query, conn, params=[pms_account_code])
    conn.close()
    return df

# Function to generate the welcome letter
def generate_welcome_letter(pms_account_code):
    # Fetch data from the database for the given PMS account code
    data_from_database = fetch_data_from_database(pms_account_code)

    # Check if data_from_database is not None before accessing its attributes
    if not data_from_database.empty:
        data_from_database['ActiveDate'] = pd.to_datetime(data_from_database['ActiveDate'])
        data = {
            "Address": data_from_database.Address.iloc[0],
            "Client Name": data_from_database.Clientname.iloc[0],
            "Date of Activation": data_from_database.ActiveDate.dt.strftime('%Y-%m-%d').iloc[0],
            "PMS Account Code": data_from_database.backofficecodeequity.iloc[0],
            "Strategy Opted": data_from_database.SchemeCodename.iloc[0],
            "Strategy Bench Mark": data_from_database.BenchMark.iloc[0],
            "Registered email id": data_from_database.EMAIL.iloc[0],
            "Registered Mobile no.": data_from_database.mobile_no.iloc[0],
            "Name of Distributor": data_from_database.IntroducerDistributorName.iloc[0],
            "Name of RM": data_from_database.RelationshipManager.iloc[0],
            "RM email id": data_from_database.RMEmail.iloc[0],
            "Mobile no.": data_from_database.MobileNo.iloc[0],
            "Total Corpus": data_from_database.TotalCorpusIntroduced.iloc[0],
            "Fund": data_from_database.Fund.iloc[0],
            "Securities": data_from_database.Securities.iloc[0],
            "Login Id": data_from_database.LoginId.iloc[0],
            "pass": data_from_database.Password.iloc[0],
            # Add more placeholders as needed
        }

        # Load the existing Word document template
        doc = Document("C:\\Users\\amograne\\Welcome Letter Folder\\maindoc.docx")  # Replace with the path to your Word document

        # Function to update placeholders in both paragraphs and tables
        def update_placeholders(doc, data):
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)  # Set the font size to 9 for each run

                for key, value in data.items():
                    placeholder = f"<<{key}>>"
                    if placeholder in paragraph.text:
                        # Replace in the entire paragraph text
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
                        for run in paragraph.runs:
                            run.font.size = Pt(9)  # Set the font size after replacement

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in data.items():
                            placeholder = f"<<{key}>>"
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))
                                for run in cell.paragraphs[0].runs:
                                    run.font.size = Pt(9)  # Set the font size after replacement

        # Update placeholders and set the font size in the document
        update_placeholders(doc, data)

        # Save the modified document
        output_path = f"welcome_letter_{pms_account_code}.docx"
        doc.save(output_path)
        return output_path
    else:
        return None

 # Web Script 

# Streamlit app title
st.title("Welcome Letter Generator")

# Placeholder for dynamic content
placeholder = st.empty()

# Input field for PMS Account Code
input_code = placeholder.text_input('Enter PMS Account Code:', key=1)

# Generate Welcome Letter button
if input_code:
    result = generate_welcome_letter(input_code)
    if result:
        # Display success message and the generated content
        st.success("Welcome letter generated successfully.")
        st.text(result)
    else:
        # Display a warning if no data is found
        st.warning("No data found for the provided PMS account code.")

    # Clear button
    click_clear = st.button('Clear Input', key=3)
    if click_clear:
        # Clear the input field by updating the value
        input_code = placeholder.text_input('Enter PMS Account Code:', value='', key=2)

        click_clear = False
